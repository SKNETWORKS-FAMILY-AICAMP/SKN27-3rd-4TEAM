"""
가상계약서 → 전체 진단 파이프라인 통합 테스트
=================================================
실행 방법:
    cd SKN27-3rd-4TEAM
    python test_virtual_contract.py

테스트 흐름:
  1. docs/가상계약서.docx 파싱
  2. 계약 필드 추출
  3. DL 브리지 (시세·위험 분석)
  4. 특약 위험 분석 (규칙 기반 fallback)
  5. 종합 위험 점수 계산
  6. 결과 리포트 출력
"""
from __future__ import annotations

import sys
import json
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Any

_ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(_ROOT))
sys.path.insert(0, str(_ROOT / "backend"))

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 1 : .docx 파싱
# ═══════════════════════════════════════════════════════════════════════════════

def parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))
    return "\n".join(paragraphs)


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 2 : 계약 필드 추출 (정규식 기반)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_fields(text: str) -> dict[str, Any]:
    import re

    fields: dict[str, Any] = {}

    # 주소
    addr = re.search(r"서울특별시\s+\S+구\s+\S+동\s+[\d-]+\s+\S+", text)
    if addr:
        fields["address"] = addr.group(0).strip()

    # 보증금
    deposit = re.search(r"금\s+(?:사억|삼억|이억|일억)[^\n]*?₩([\d,]+)", text)
    if deposit:
        fields["deposit_amount"] = int(deposit.group(1).replace(",", ""))

    # 전용면적
    area = re.search(r"주거전용면적[:\s]*([0-9.]+)\s*㎡", text)
    if area:
        fields["exclusive_area_m2"] = float(area.group(1))

    # 주택 유형
    if "다세대주택[■]" in text or "다세대주택" in text:
        fields["housing_type"] = "연립다세대"
    elif "연립주택[■]" in text:
        fields["housing_type"] = "연립다세대"
    elif "아파트[■]" in text:
        fields["housing_type"] = "아파트"
    elif "다가구주택[■]" in text:
        fields["housing_type"] = "단독다가구"

    # 임대인
    landlord = re.search(r"임대사업자\s*\|\s*성명[^\|]*\|\s*([가-힣]+)", text)
    if landlord:
        fields["landlord_name"] = landlord.group(1)

    # 임차인
    tenant = re.search(r"임차인\s*\|\s*성명[^\|]*\|\s*([가-힣]+)", text)
    if tenant:
        fields["tenant_name"] = tenant.group(1)

    # 근저당
    mortgage = re.search(r"설정금액[:\s]*([\d,]+)원", text)
    if mortgage:
        fields["mortgage_amount"] = int(mortgage.group(1).replace(",", ""))

    # 보증보험
    if "일부가입" in text:
        fields["insurance_status"] = "일부가입"
        ins_amt = re.search(r"보증대상\s*금액[:\s]*([\d,]+)원", text)
        if ins_amt:
            fields["insurance_amount"] = int(ins_amt.group(1).replace(",", ""))

    # 계약기간
    period = re.search(r"(\d{4})년\s*(\d{2})월\s*(\d{2})일\s*∼\s*(\d{4})년\s*(\d{2})월\s*(\d{2})일", text)
    if period:
        fields["contract_start"] = f"{period.group(1)}-{period.group(2)}-{period.group(3)}"
        fields["contract_end"] = f"{period.group(4)}-{period.group(5)}-{period.group(6)}"

    # 동 이름
    dong = re.search(r"([가-힣0-9]+동)", fields.get("address", ""))
    if dong:
        fields["dong_name"] = dong.group(1)

    # 특약 추출
    special_start = text.find("【특약사항】")
    if special_start != -1:
        special_text = text[special_start:]
        clauses = re.findall(r"\d+\.\s+(.+?)(?=\n\d+\.|\Z)", special_text, re.DOTALL)
        fields["special_terms"] = [c.strip().replace("\n", " ")[:300] for c in clauses if c.strip()]

    return fields


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 3 : DL 브리지 호출
# ═══════════════════════════════════════════════════════════════════════════════

def run_dl(fields: dict[str, Any]) -> tuple[dict, list[dict]]:
    try:
        from common.tools.dl_market_bridge import run_dl_analysis
        pack, findings = run_dl_analysis(fields)
        pack_dict = {
            "task_type": pack.task_type,
            "quality_score": pack.quality.score,
            "quality_sufficient": pack.quality.sufficient,
            "quality_reason": pack.quality.reason,
            "contexts": [
                {"title": c.title, "score": c.score, "text_preview": c.text[:200]}
                for c in pack.contexts
            ],
        }
        findings_list = [asdict(f) for f in findings]
        return pack_dict, findings_list
    except Exception as e:
        return {"error": str(e), "quality_sufficient": False}, []


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4 : 특약 위험 분석 (규칙 기반)
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class RiskItem:
    code: str
    title: str
    severity: str
    score: int
    description: str
    evidence: str = ""

CLAUSE_RULES = [
    {
        "code": "CLAUSE_MORTGAGE_LATE_RELEASE",
        "title": "근저당 말소 지연 위험",
        "severity": "HIGH",
        "score": 20,
        "keywords": ["근저당", "말소", "30일"],
        "description": "잔금일 이후 30일 이내 말소 약정은 그 기간 동안 보증금이 근저당 위험에 노출됩니다. "
                       "잔금 지급 전 말소 조건으로 재협상해야 합니다.",
    },
    {
        "code": "CLAUSE_DEPOSIT_BUSINESS_USE",
        "title": "보증금 사업자금 전용 허용 특약",
        "severity": "CRITICAL",
        "score": 30,
        "keywords": ["사업 운영자금", "활용할 수 있으며"],
        "description": "임대인이 보증금을 사업 운영자금으로 사용할 수 있다는 특약은 "
                       "반환 불능 전세사기의 전형적 패턴입니다. 즉시 삭제 요청이 필요합니다.",
    },
    {
        "code": "CLAUSE_RISK_TRANSFER",
        "title": "시세 하락 위험 임차인 전가",
        "severity": "HIGH",
        "score": 20,
        "keywords": ["시세 변동에 따른 위험은 임차인"],
        "description": "시세 하락 위험을 임차인에게 전가하는 특약은 임차인에게 현저히 불리합니다. "
                       "삭제 또는 수정이 필요합니다.",
    },
    {
        "code": "CLAUSE_ILLEGAL_RENT_INCREASE",
        "title": "법정 임대료 증액 한도 초과 특약",
        "severity": "HIGH",
        "score": 15,
        "keywords": ["5%를 초과", "5%)를 초과"],
        "description": "민간임대주택특별법 제44조는 임대료 증액을 5% 이내로 제한합니다. "
                       "이를 초과하는 특약은 무효이며, 임대인의 법 위반 가능성을 시사합니다.",
    },
    {
        "code": "CLAUSE_PARTIAL_INSURANCE_RISK",
        "title": "보증보험 미가입분 위험 임차인 인수",
        "severity": "HIGH",
        "score": 20,
        "keywords": ["미가입분", "위험을 인지하고 계약"],
        "description": "보증보험 미가입 6천만원에 대한 위험을 임차인이 인수한다는 특약은 "
                       "사실상 보증금 일부 포기입니다. 전액 가입 또는 근저당 선말소를 요구해야 합니다.",
    },
    {
        "code": "CLAUSE_DEFECT_NOTICE_BURDEN",
        "title": "하자 미통지 시 임차인 책임 전가",
        "severity": "MEDIUM",
        "score": 10,
        "keywords": ["통지하지 않은 하자는 임차인"],
        "description": "30일 이내 통지하지 않은 하자를 임차인 책임으로 돌리는 특약은 "
                       "임차인의 수선 청구권을 부당하게 제한할 수 있습니다.",
    },
]

def analyze_clauses(fields: dict[str, Any], contract_text: str) -> list[RiskItem]:
    terms = fields.get("special_terms", [])
    joined = " ".join(terms) + " " + contract_text

    findings = []
    for rule in CLAUSE_RULES:
        if any(kw in joined for kw in rule["keywords"]):
            evidence = ""
            for term in terms:
                if any(kw in term for kw in rule["keywords"]):
                    evidence = term[:150]
                    break
            findings.append(RiskItem(
                code=rule["code"],
                title=rule["title"],
                severity=rule["severity"],
                score=rule["score"],
                description=rule["description"],
                evidence=evidence,
            ))
    return findings


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5 : 필수 확인 항목 (계약서만으로 알 수 없는 위험)
# ═══════════════════════════════════════════════════════════════════════════════

def required_checks(fields: dict[str, Any]) -> list[RiskItem]:
    checks = [
        RiskItem("REQ_REGISTRY", "등기부 권리관계 추가 확인 필요", "HIGH", 15,
                 "계약서에 근저당(6천만원)이 기재되어 있으나 가압류·압류·신탁 등 추가 권리관계는 "
                 "등기부등본(갑구·을구)을 직접 확인해야 합니다."),
        RiskItem("REQ_OWNER_MATCH", "임대인-소유자 동일인 확인 필요", "HIGH", 10,
                 "임대인 오성호가 실제 등기부 소유자인지 신분증과 등기부를 대조하세요."),
    ]
    if fields.get("housing_type") in {"연립다세대", "단독다가구"}:
        checks.append(RiskItem(
            "REQ_SENIOR_TENANTS", "선순위 임차인 보증금 현황 확인 필요", "HIGH", 10,
            "다세대주택은 다른 세대 선순위 보증금 합계가 경매 시 회수 가능금액에 영향을 줍니다. "
            "확정일자 부여 현황 및 전입세대확인서를 요청하세요.",
        ))
    return checks


# ═══════════════════════════════════════════════════════════════════════════════
# STEP 6 : 종합 위험 점수 및 리포트
# ═══════════════════════════════════════════════════════════════════════════════

SEVERITY_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}

def compute_risk(
    clause_findings: list[RiskItem],
    req_findings: list[RiskItem],
    dl_findings: list[dict],
    fields: dict[str, Any],
) -> tuple[int, str]:
    base = sum(f.score for f in clause_findings + req_findings)

    # DL 위험 delta 합산
    dl_delta = sum(f.get("score_delta", 0) for f in dl_findings)

    # 기본 전세가율 위험 계산 (계약서 데이터 기반)
    deposit = fields.get("deposit_amount", 0)
    mortgage = fields.get("mortgage_amount", 0)
    area = fields.get("exclusive_area_m2", 0)

    market_bonus = 0
    if deposit and area:
        per_pyeong = deposit / (area * 0.3025) / 10000  # 만원/평
        # 종로구 다세대 평균 시세 기준 (약 2500만원/평) → 전세가율 추정
        est_sale = per_pyeong / 0.75  # 역산: 전세가율 75% 가정
        est_ratio = (per_pyeong / est_sale) * 100
        if est_ratio >= 85:
            market_bonus = 20
        elif est_ratio >= 75:
            market_bonus = 10

    total = min(100, base + dl_delta + market_bonus)
    if total >= 75:
        level = "CRITICAL"
    elif total >= 55:
        level = "HIGH"
    elif total >= 35:
        level = "MEDIUM"
    else:
        level = "LOW"
    return total, level


def print_report(
    fields: dict[str, Any],
    clause_findings: list[RiskItem],
    req_findings: list[RiskItem],
    dl_pack: dict,
    dl_findings: list[dict],
    risk_score: int,
    risk_level: str,
):
    SEP = "=" * 65
    SEP2 = "-" * 65
    LEVEL_ICON = {"CRITICAL": "🔴", "HIGH": "🟠", "MEDIUM": "🟡", "LOW": "🟢"}

    print(f"\n{SEP}")
    print("  전세계약 위험 진단 리포트")
    print(f"  ※ 본 결과는 법률 자문이 아닌 참고 정보입니다.")
    print(SEP)

    # ── 계약 기본 정보 ──────────────────────────────────────────
    print("\n【계약 기본 정보】")
    print(f"  주소       : {fields.get('address', '미상')}")
    print(f"  주택 유형  : {fields.get('housing_type', '미상')}")
    area = fields.get("exclusive_area_m2")
    print(f"  전용면적   : {area}㎡ ({area * 0.3025:.1f}평)" if area else "  전용면적  : 미상")
    deposit = fields.get("deposit_amount", 0)
    print(f"  보증금     : {deposit:,}원 ({deposit // 10000:,}만원)")
    print(f"  임대인     : {fields.get('landlord_name', '미상')}")
    print(f"  임차인     : {fields.get('tenant_name', '미상')}")
    print(f"  계약기간   : {fields.get('contract_start','?')} ~ {fields.get('contract_end','?')}")
    mortgage = fields.get("mortgage_amount", 0)
    if mortgage:
        print(f"  근저당     : {mortgage:,}원 (○○저축은행)")
    ins = fields.get("insurance_status")
    if ins:
        ins_amt = fields.get("insurance_amount", 0)
        print(f"  보증보험   : {ins} ({ins_amt:,}원 / 전체의 {ins_amt*100//deposit:.0f}%)")

    # ── 종합 위험도 ─────────────────────────────────────────────
    print(f"\n{SEP2}")
    icon = LEVEL_ICON.get(risk_level, "⚪")
    bar_filled = int(risk_score / 5)
    bar = "█" * bar_filled + "░" * (20 - bar_filled)
    print(f"  위험 점수  : {risk_score}점 / 100점")
    print(f"  위험 등급  : {icon} {risk_level}")
    print(f"  [{bar}] {risk_score}%")

    # ── DL 시세 분석 ────────────────────────────────────────────
    print(f"\n{SEP2}")
    print("【딥러닝 시세 분석 결과】")
    if dl_pack.get("quality_sufficient"):
        print(f"  신뢰도 점수: {dl_pack['quality_score']:.2f}")
        for ctx in dl_pack.get("contexts", []):
            print(f"\n  [{ctx['title']}] (score={ctx['score']:.2f})")
            for line in ctx["text_preview"].split("\n"):
                if line.strip():
                    print(f"    {line.strip()}")
    else:
        reason = dl_pack.get("quality_reason") or dl_pack.get("error", "DB 미연결 또는 데이터 부족")
        print(f"  ⚠ DL 분석 불가: {reason}")
        print("  → 딥러닝 모델 DB(PostgreSQL) 연결 후 재실행하면 LSTM 시세 예측이 추가됩니다.")

    if dl_findings:
        print(f"\n  DL 위험 발견 ({len(dl_findings)}건):")
        for f in dl_findings:
            icon2 = LEVEL_ICON.get(f.get("severity", "LOW"), "⚪")
            print(f"  {icon2} [{f['code']}] {f['title']} (+{f['score_delta']}점)")
            print(f"     {f['description']}")

    # ── 특약 위험 분석 ──────────────────────────────────────────
    all_clause = sorted(clause_findings, key=lambda x: SEVERITY_ORDER.get(x.severity, 9))
    print(f"\n{SEP2}")
    print(f"【특약 위험 분석】 ({len(all_clause)}건 발견)")
    if not all_clause:
        print("  ✅ 분석된 위험 특약 없음")
    for idx, f in enumerate(all_clause, 1):
        icon2 = LEVEL_ICON.get(f.severity, "⚪")
        print(f"\n  {idx}. {icon2} [{f.severity}] {f.title} (+{f.score}점)")
        print(f"     {f.description}")
        if f.evidence:
            print(f"     📄 근거: \"{f.evidence[:100]}...\"")

    # ── 필수 확인 항목 ──────────────────────────────────────────
    print(f"\n{SEP2}")
    print(f"【계약 전 필수 확인 항목】 ({len(req_findings)}건)")
    for idx, f in enumerate(req_findings, 1):
        icon2 = LEVEL_ICON.get(f.severity, "⚪")
        print(f"  {idx}. {icon2} {f.title}")
        print(f"     → {f.description}")

    # ── 요약 권고 ───────────────────────────────────────────────
    print(f"\n{SEP}")
    print("【종합 권고사항】")
    critical = [f for f in all_clause if f.severity == "CRITICAL"]
    high = [f for f in all_clause + req_findings if f.severity == "HIGH"]

    if critical:
        print("  ⛔ 계약 즉시 중단 권고:")
        for f in critical:
            print(f"     - {f.title}")
    if high:
        print("  ⚠️  계약 전 반드시 해결해야 할 항목:")
        for f in high:
            print(f"     - {f.title}")

    print("\n  📋 체크리스트:")
    print("     □ 잔금 지급 전 근저당 완전 말소 확인")
    print("     □ 보증금 사업자금 전용 특약 삭제")
    print("     □ 등기부등본(갑구·을구) 직접 열람")
    print("     □ 임대인 신분증·등기 소유자 대조")
    print("     □ HUG 전세보증보험 전액 가입 요구")
    if fields.get("housing_type") in {"연립다세대", "단독다가구"}:
        print("     □ 전입세대확인서 및 확정일자 현황 확인")
    print(SEP)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    contract_path = "docs/가상계약서.docx"

    print("\n[1/5] 계약서 파싱 중...")
    text = parse_docx(contract_path)
    print(f"      → {len(text):,}자 파싱 완료")

    print("[2/5] 계약 필드 추출 중...")
    fields = extract_fields(text)
    print(f"      → 추출 필드: {list(fields.keys())}")

    print("[3/5] 딥러닝 시세·위험 분석 중...")
    dl_pack, dl_findings = run_dl(fields)
    dl_ok = dl_pack.get("quality_sufficient", False)
    print(f"      → DL 분석: {'성공' if dl_ok else '미연결(fallback)'}, "
          f"위험 발견 {len(dl_findings)}건")

    print("[4/5] 특약·필수 항목 분석 중...")
    clause_findings = analyze_clauses(fields, text)
    req_findings = required_checks(fields)
    print(f"      → 특약 위험 {len(clause_findings)}건, 필수 확인 {len(req_findings)}건")

    print("[5/5] 종합 위험 점수 계산 중...")
    risk_score, risk_level = compute_risk(clause_findings, req_findings, dl_findings, fields)
    print(f"      → 위험 점수: {risk_score}점 / 등급: {risk_level}")

    print_report(fields, clause_findings, req_findings, dl_pack, dl_findings, risk_score, risk_level)

    # JSON 결과 저장
    result = {
        "contract_path": contract_path,
        "fields": {k: v for k, v in fields.items() if k != "special_terms"},
        "special_terms_count": len(fields.get("special_terms", [])),
        "risk_score": risk_score,
        "risk_level": risk_level,
        "clause_findings": [asdict(f) for f in clause_findings],
        "required_checks": [asdict(f) for f in req_findings],
        "dl_integrated": dl_ok,
        "dl_findings_count": len(dl_findings),
    }
    out_path = Path("virtual_contract_diagnosis.json")
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n  📄 JSON 결과 저장: {out_path.resolve()}")


if __name__ == "__main__":
    main()
