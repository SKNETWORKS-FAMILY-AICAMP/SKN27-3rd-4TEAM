"""Contract document parsing and lightweight field extraction."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from langchain_core.tools import tool

from common.tools.llm import LLMUnavailable, extract_json_object, ollama_generate

MOCK_CONTRACT_TEXT = """
전세계약서
임대인 홍길동 임차인 김철수
소재지 서울특별시 종로구 청운동 12-3
보증금 금 200,000,000원 월세 0원
계약기간 2025.03.01부터 2027.03.01까지
주택유형 연립다세대 전용면적 52.3㎡
특약사항
1. 임차인은 시설물 수리비를 전액 부담한다.
2. 임대인은 잔금일 다음날까지 근저당권 등 권리변동을 하지 않는다.
""".strip()


def parse_contract_file(file_path: str | None) -> tuple[str, list[str], float | None]:
    if not file_path:
        return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None

    path = Path(file_path)
    if not path.exists():
        return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return text, [text], None

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            # 단락 텍스트 추출 (빈 단락 제외)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # 표 안의 텍스트도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        paragraphs.append(" | ".join(row_texts))
            text = "\n".join(paragraphs)
            if text.strip():
                # docx는 페이지 구분이 없으므로 전체를 하나의 페이지로 처리
                return text, [text], None
        except Exception as e:
            print(f"[document] docx 파싱 실패: {e}")
            return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None

        return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
            text = "\n\n".join(page for page in pages if page)
            if _looks_readable_korean(text):
                return text, pages, None
        except Exception:
            pass

        try:
            import fitz

            document = fitz.open(str(path))
            pages = [page.get_text("text").strip() for page in document]
            document.close()
            text = "\n\n".join(page for page in pages if page)
            if text:
                return text, pages, None
        except Exception:
            return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None

        return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None

    return MOCK_CONTRACT_TEXT, [MOCK_CONTRACT_TEXT], None


def extract_contract_fields(text: str) -> dict[str, Any]:
    try:
        fields = _extract_with_llm(text)
        fields["extraction_method"] = "ollama"
        return fields
    except Exception:
        fields = _extract_with_regex(text)
        fields["extraction_method"] = "regex_fallback"
        return fields


def _extract_with_llm(text: str) -> dict[str, Any]:
    system = "너는 전세계약서에서 구조화 정보를 추출하는 한국어 정보추출 에이전트다. JSON만 반환한다."
    prompt = f"""
다음 전세계약서 텍스트에서 정보를 추출해 JSON으로만 반환해.
숫자는 만원 단위 정수로 정규화해. 모르면 null 또는 빈 배열을 사용해.

필드:
landlord, tenant, address, dong_name, deposit_amount, monthly_rent,
contract_start, contract_end, housing_type, exclusive_area_m2, special_terms

계약서:
{text[:6000]}
""".strip()
    raw = ollama_generate(prompt, system=system)
    data = extract_json_object(raw)
    if not isinstance(data, dict):
        raise LLMUnavailable("LLM returned non-object JSON")
    return data


def _extract_with_regex(text: str) -> dict[str, Any]:
    normalized = re.sub(r"[ \t]+", " ", text)
    deposit = _money_to_manwon(_first_match(normalized, [r"보증금\s*(?:금)?\s*([0-9,]+)\s*원", r"보증금\s*([0-9,]+)\s*만원"]))
    rent = _money_to_manwon(_first_match(normalized, [r"월세\s*(?:금)?\s*([0-9,]+)\s*원", r"월세\s*([0-9,]+)\s*만원"]))
    area_raw = _first_match(normalized, [r"전용면적\s*([0-9.]+)\s*(?:㎡|m2)", r"면적\s*([0-9.]+)\s*(?:㎡|m2)"])
    address = _first_match(normalized, [r"소재지\s*([^\n]+)", r"주소\s*([^\n]+)"])
    housing_type = _first_match(normalized, [r"(연립다세대|오피스텔|단독다가구|단독|다가구|아파트)"])
    terms = _extract_special_terms(text)

    return {
        "landlord": _first_match(normalized, [r"임대인\s*([가-힣A-Za-z]{2,10})"]),
        "tenant": _first_match(normalized, [r"임차인\s*([가-힣A-Za-z]{2,10})"]),
        "address": address,
        "dong_name": _extract_dong(address or normalized),
        "deposit_amount": deposit,
        "monthly_rent": rent,
        "contract_start": _first_match(normalized, [r"([0-9]{4}[.\-/][0-9]{1,2}[.\-/][0-9]{1,2})\s*부터"]),
        "contract_end": _first_match(normalized, [r"부터\s*([0-9]{4}[.\-/][0-9]{1,2}[.\-/][0-9]{1,2})\s*까지"]),
        "housing_type": _normalize_housing_type(housing_type),
        "exclusive_area_m2": float(area_raw) if area_raw else None,
        "special_terms": terms,
    }


def _looks_readable_korean(text: str) -> bool:
    if not text.strip():
        return False
    korean_chars = len(re.findall(r"[가-힣]", text))
    replacement_chars = text.count("?") + text.count("�")
    return korean_chars >= 10 and replacement_chars < max(10, len(text) * 0.1)
def _first_match(text: str, patterns: list[str]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    return None


def _money_to_manwon(value: str | None) -> int | None:
    if not value:
        return None
    number = int(value.replace(",", ""))
    return number // 10000 if number > 100000 else number


def _extract_dong(text: str) -> str | None:
    match = re.search(r"([가-힣0-9]+동)", text)
    return match.group(1) if match else None


def _normalize_housing_type(value: str | None) -> str | None:
    if not value:
        return None
    if value in {"단독", "다가구", "단독다가구"}:
        return "단독다가구"
    if value == "연립다세대":
        return "연립다세대"
    return value


def _extract_special_terms(text: str) -> list[str]:
    marker = re.search(r"특약(?:사항)?", text)
    if not marker:
        return []
    tail = text[marker.end():]
    lines = [line.strip(" -\t") for line in tail.splitlines()]
    return [line for line in lines if line and len(line) > 4][:20]




@tool
def parse_contract_file_tool(file_path: str | None = None) -> tuple[str, list[str], float | None]:
    """Parse a contract PDF/TXT path and return text, page texts, and OCR confidence."""
    return parse_contract_file(file_path)


@tool
def extract_contract_fields_tool(text: str) -> dict[str, Any]:
    """Extract structured jeonse contract fields from contract text."""
    return extract_contract_fields(text)
