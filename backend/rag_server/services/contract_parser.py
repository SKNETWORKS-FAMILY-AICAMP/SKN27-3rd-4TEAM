"""
전세계약 위험 진단 에이전트 - 전세계약서 파서
"""
from __future__ import annotations
import re, io, json, os
import pdfplumber
from rag_server.models.schemas import ContractInfo

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    from PIL import ImageFilter, ImageEnhance
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

RISK_KEYWORDS = [
    "전세가율","시세","매매가",
    "근저당","가압류","가처분","저당권","담보","선순위","후순위",
    "미등기","무허가","건축물대장",
    "현금","계좌","확정일자","전입신고","대항력","임차권등기",
    "대리인","위임","인감","소유자",
    "원상복구","수리","하자","누수","도배","장판",
    "전세보증보험","HUG","SGI",
    "보증금","계약금","잔금",
]

CONTRACT_EXTRACT_PROMPT = """아래는 한국 부동산 임대차 계약서 텍스트입니다.
계약서 종류(일반 전세·월세, 공인중개사 표준, 민간임대, 법인계약 등)와 무관하게 아래 항목을 추출하여 JSON만 출력하세요.
찾을 수 없으면 반드시 JSON null을 사용하세요(문자열 "null" 금지).

[계약서 텍스트]
{text}

[추출 규칙]
- lessor_name: 집을 빌려주는 측(임대인/임대사업자/소유자/갑)의 실제 한글 성명(2~5자).
  역할명("임대인","임대사업자","소유자","갑","을","또는","법인" 등) 자체는 반환 금지.
  성명란이 여럿이면 임대인 측 성명만 추출. 없으면 null.

- lessee_name: 집을 빌리는 측(임차인/세입자/을)의 실제 한글 성명(2~5자).
  역할명 자체는 반환 금지. "성명:" "이름:" 뒤 이름, 서명·날인란의 이름 등 다양한 위치를 탐색.
  없으면 null.

- address: 임대 목적물(계약 대상 부동산)의 주소 전체 문자열.
  "소재지", "주소", "목적물", "임대차 목적물", "주택 소재지", "부동산의 표시" 등 다양한 표현 뒤 주소를 탐색.
  없으면 null.

- deposit_amount: 임대보증금(전세금) 총액(만원 단위 정수).
  한글 금액("사억원정"→40000), 숫자("400,000,000"→40000), 혼합("4억"→40000) 모두 처리.
  분할 납부액(계약금·중도금·잔금)의 합산이 아닌 총액 기재값을 우선.
  보험·담보 관련 부분 금액은 제외. 없으면 null.

- monthly_rent: 월세·월임대료(만원 단위 정수). 전세(보증금만)이거나 기재 없으면 0.
  한글 금액("영원정"→0) 포함 처리. 없으면 0.

- contract_start: 임대차(거주) 기간 시작일 YYYY-MM-DD.
  계약 체결일(서명일)과 구별하여 실제 입주·거주 시작일을 추출.
  "계약기간", "임대차기간", "기간" 뒤 범위의 첫 날짜. 없으면 null.

- contract_end: 임대차(거주) 기간 종료일 YYYY-MM-DD.
  같은 범위의 마지막 날짜. 계약기간과 무관한 의무기간·지정기간 종료일과 혼동 금지.
  없으면 null.

- special_terms: "특약사항" 또는 이에 준하는 별도 합의 조항 섹션의 내용 전체.
  표준 조항(본문 조항)은 제외하고 당사자 간 별도로 합의한 조항만 추출.
  섹션 제목(특약사항)은 포함하지 않고 내용만 반환. 없으면 null.
  텍스트에 "[특약사항]" 태그가 있으면 해당 태그 이하 내용을 우선 사용.

[출력 형식]
{{
  "lessor_name": null,
  "lessee_name": null,
  "address": null,
  "deposit_amount": null,
  "monthly_rent": 0,
  "contract_start": null,
  "contract_end": null,
  "special_terms": null
}}"""


class ContractParser:

    @classmethod
    def from_text(cls, text: str) -> ContractInfo:
        return cls()._parse(text)

    @classmethod
    def from_pdf_bytes(cls, pdf_bytes: bytes) -> ContractInfo:
        return cls.from_text(cls._extract_pdf_text(pdf_bytes))

    @classmethod
    def from_docx_bytes(cls, docx_bytes: bytes) -> ContractInfo:
        return cls.from_text(cls._extract_docx_text(docx_bytes))

    @staticmethod
    def _extract_docx_text(docx_bytes: bytes) -> str:
        if not DOCX_AVAILABLE:
            return "[docx 파싱 오류: python-docx 라이브러리 없음]"
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))

            # 특약사항 표 먼저 탐색:
            # - 1개 셀짜리 표이고
            # - 내용이 "1. " 로 시작하는 번호 조항 블록이며
            # - 개인정보 관련 텍스트가 아닌 것
            special_terms_text = None
            for table in doc.tables:
                if len(table.rows) == 1 and len(table.columns) == 1:
                    cell_text = table.rows[0].cells[0].text.strip()
                    if (re.match(r"1\.\s+", cell_text)
                            and "개인정보" not in cell_text[:30]
                            and len(cell_text) > 50):
                        special_terms_text = cell_text
                        print(f"[ContractParser] docx 특약사항 표 발견: {len(cell_text)}자")
                        break

            # 본문 단락 수집 (빈 줄 제외)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # 일반 표 데이터 수집 (특약사항 표는 제외)
            table_rows = []
            for table in doc.tables:
                # 특약사항 표(1x1, 번호조항)는 건너뜀
                if (len(table.rows) == 1 and len(table.columns) == 1
                        and table.rows[0].cells[0].text.strip() == (special_terms_text or "")):
                    continue
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    unique_cells = list(dict.fromkeys(c for c in cells if c))
                    if unique_cells:
                        table_rows.append(" | ".join(unique_cells))

            body = "\n".join(paragraphs)
            table_text = "\n".join(table_rows)

            result = body
            if table_text:
                result += "\n\n[표 데이터]\n" + table_text
            if special_terms_text:
                result += "\n\n[특약사항]\n" + special_terms_text

            print(f"[ContractParser] docx 추출: 단락 {len(paragraphs)}개, 표 행 {len(table_rows)}개"
                  + (f", 특약사항 {len(special_terms_text)}자" if special_terms_text else ""))
            return result
        except Exception as e:
            print(f"[ContractParser] docx 추출 오류: {e}")
            return f"[docx 파싱 오류: {e}]"

    @staticmethod
    def _extract_pdf_text(pdf_bytes: bytes) -> str:
        parts, table_parts = [], []
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or "")
                    try:
                        for table in (page.extract_tables() or []):
                            for row in table:
                                if not row:
                                    continue
                                row_text = " | ".join((c or "").strip() for c in row if c)
                                if row_text.strip():
                                    table_parts.append(row_text)
                    except Exception:
                        pass
        except Exception as e:
            return f"[PDF 추출 오류: {e}]"

        body = "\n".join(parts).strip()
        table = "\n".join(table_parts).strip()

        if table:
            print(f"[ContractParser] pdfplumber 표 추출: {len(table_parts)}개 행")
            text = body + "\n\n[표 데이터]\n" + table
        else:
            text = body

        if not text.strip() and OCR_AVAILABLE:
            print("[ContractParser] 텍스트 추출 실패 -> OCR 시도 중...")
            try:
                images = convert_from_bytes(pdf_bytes, dpi=400)
                ocr_parts = []
                for i, image in enumerate(images):
                    image = image.convert("L")
                    image = ImageEnhance.Contrast(image).enhance(2.0)
                    image = image.filter(ImageFilter.SHARPEN)
                    image = image.point(lambda p: 255 if p > 128 else 0)
                    ocr_parts.append(pytesseract.image_to_string(
                        image, lang="kor+eng", config="--psm 6 --oem 1"))
                    print(f"  [OCR] {i+1}/{len(images)} 페이지 완료")
                text = "\n".join(ocr_parts).strip()
                print(f"[ContractParser] OCR 완료: {len(text)}자")
            except Exception as e:
                print(f"[ContractParser] OCR 실패: {e}")
        elif not text.strip():
            print("[ContractParser] 스캔 PDF 감지됐으나 OCR 라이브러리 없음")

        return text

    # ─────────────────────────────────────────────────────────
    def _parse(self, text: str) -> ContractInfo:
        if not text or len(text.strip()) < 20:
            return ContractInfo(raw_text=text)

        llm_info = self._extract_with_llm(text)

        if llm_info is None:
            print("[ContractParser] LLM 추출 실패 -> 정규식 fallback")
            return ContractInfo(
                lessor_name=self._extract_lessor(text),
                lessee_name=self._extract_lessee(text),
                address=self._extract_address(text),
                deposit_amount=self._extract_deposit(text),
                monthly_rent=self._extract_monthly_rent(text),
                contract_start=self._extract_date(text, "start"),
                contract_end=self._extract_date(text, "end"),
                special_terms=self._extract_special_terms(text),
                raw_text=text,
            )

        print("[ContractParser] LLM 추출 후 검증 및 정규식 보완")

        _ROLE = {"임대사업자","임대인","임차인","소유자","대리인","법인","성명","또는"}
        if not llm_info.lessor_name or llm_info.lessor_name.strip() in _ROLE:
            llm_info.lessor_name = self._extract_lessor(text)
            print(f"[ContractParser] lessor_name 재추출: {llm_info.lessor_name}")
        if not llm_info.lessee_name or llm_info.lessee_name.strip() in _ROLE:
            llm_info.lessee_name = self._extract_lessee(text)
        if not llm_info.address:
            llm_info.address = self._extract_address(text)

        regex_dep = self._extract_deposit(text)
        if regex_dep and (not llm_info.deposit_amount or regex_dep > llm_info.deposit_amount):
            llm_info.deposit_amount = regex_dep
            print(f"[ContractParser] deposit_amount 정규식 우선: {regex_dep}만원")
        if llm_info.monthly_rent is None:
            llm_info.monthly_rent = self._extract_monthly_rent(text) or 0

        # ── 날짜 교차 검증 (cascading error 방지)
        try:
            from datetime import date as _dt
            today = _dt.today()
            if llm_info.contract_start:
                cs = _dt.fromisoformat(llm_info.contract_start)
                if (today - cs).days > 3 * 365:
                    print(f"[ContractParser] contract_start 오탐({llm_info.contract_start}, 3년 초과) -> null")
                    llm_info.contract_start = None
                    llm_info.contract_end = None
            if llm_info.contract_start and llm_info.contract_end:
                cs = _dt.fromisoformat(llm_info.contract_start)
                ce = _dt.fromisoformat(llm_info.contract_end)
                if (ce - cs).days < 365:
                    print(f"[ContractParser] contract_end 오탐({(ce-cs).days}일 < 365) -> null")
                    llm_info.contract_end = None
        except (ValueError, TypeError):
            llm_info.contract_start = None
            llm_info.contract_end = None

        if not llm_info.contract_start:
            llm_info.contract_start = self._extract_date(text, "start")
        if not llm_info.contract_end:
            llm_info.contract_end = self._extract_contract_end(text, llm_info.contract_start)
        if not llm_info.special_terms:
            llm_info.special_terms = self._extract_special_terms(text)

        return llm_info

    # ─────────────────────────────────────────────────────────
    @staticmethod
    def _extract_with_llm(text: str) -> ContractInfo | None:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

            # 표 데이터 분리 (항상 LLM 프롬프트에 포함)
            if "\n\n[표 데이터]\n" in text:
                body, raw_table = text.split("\n\n[표 데이터]\n", 1)
                table_sec = "\n\n[표 데이터 - 이름/금액/날짜가 표 셀에 있을 수 있음]\n" + raw_table[:1200]
            else:
                body, table_sec = text, ""

            # 특약사항 섹션 추출 (다양한 계약서 구조 대응)
            special_sec = ""
            # docx에서 추출된 [특약사항] 태그가 있으면 바로 사용
            if "\n\n[특약사항]\n" in text:
                _, _st_raw = text.split("\n\n[특약사항]\n", 1)
                special_sec = "\n\n[특약사항 섹션]\n" + _st_raw.strip()[:1500]
            else:
                _st = ContractParser._find_special_terms_text(text)
                if _st:
                    special_sec = "\n\n[특약사항 섹션]\n" + _st[:1500]

            combined = body[:3200] + table_sec + special_sec
            prompt = CONTRACT_EXTRACT_PROMPT.format(text=combined)

            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0, max_tokens=1000)
            raw = resp.choices[0].message.content.strip()

            m = re.search(r"\{[\s\S]+\}", raw)
            if not m:
                return None
            data = json.loads(m.group(0))

            def _clean(v):
                if v is None:
                    return None
                if isinstance(v, str) and v.strip().lower() in ("null","none","n/a","없음",""):
                    return None
                return v

            def _iso(v):
                if not v:
                    return None
                v = str(v).strip()
                if re.match(r"^\d{4}-\d{2}-\d{2}$", v):
                    return v
                km = re.search(r"(\d{4})\s*년\s*(\d{1,2})\s*월\s*(\d{1,2})", v)
                return f"{km.group(1)}-{km.group(2).zfill(2)}-{km.group(3).zfill(2)}" if km else None

            dep = _clean(data.get("deposit_amount"))
            if isinstance(dep, str):
                d = re.sub(r"[^\d]", "", dep)
                dep = int(d) if d else None

            mon = _clean(data.get("monthly_rent"))
            if isinstance(mon, str):
                d = re.sub(r"[^\d]", "", mon)
                mon = int(d) if d else 0

            sp_terms = _clean(data.get("special_terms"))
            if sp_terms:
                sp_terms = re.sub(r"\n+_+\s*$", "", sp_terms).strip()
                sp_terms = re.sub(r"\b[A-Z]{2,4}\b", "", sp_terms)
                sp_terms = re.sub(r"\s{2,}", " ", sp_terms).strip()

            print("[ContractParser] LLM 추출 성공")
            return ContractInfo(
                lessor_name=_clean(data.get("lessor_name")),
                lessee_name=_clean(data.get("lessee_name")),
                address=_clean(data.get("address")),
                deposit_amount=dep,
                monthly_rent=mon or 0,
                contract_start=_iso(data.get("contract_start")),
                contract_end=_iso(data.get("contract_end")),
                special_terms=sp_terms,
                raw_text=text,
            )
        except Exception as e:
            print(f"[ContractParser] LLM 추출 오류: {e}")
            return None

    # ─────────────────────────────────────────────────────────
    @staticmethod
    def _extract_lessor(text: str) -> str | None:
        _NOT = {
            "임대인","임차인","임대사업자","성명","법인명","사업자","소유자",
            "대리인","또는","이하","에게","으로","으로서","하며","하고",
            "있으며","있으나","없으며","해야","한다","된다","한후","경우",
            "서명","날인","보관","작성","없음","해당","없이","까지","부터",
        }
        def _ok(s):
            s = s.strip()
            return 2 <= len(s) <= 5 and re.fullmatch(r"[가-힣]{2,5}", s) and s not in _NOT

        for pat in [
            r"임\s*대\s*사\s*업\s*자\s*\|[^\|]*성\s*명[^\|]*\|\s*([가-힣]{2,5})",
            r"임\s*대\s*사\s*업\s*자\s*\|[^\|]+\|\s*([가-힣]{2,5})\s*\|",
            r"임\s*대\s*사\s*업\s*자\s*성\s*명[:\s]+([가-힣]{2,5})",
            r"임\s*대\s*인\s*성\s*명[:\s]+([가-힣]{2,5})",
            r"예\s*금\s*주[:\s]*([가-힣]{2,5})",
            r"임\s*대\s*인[\s\S]{0,60}성\s*명\s*\(?\s*법\s*인\s*명\s*\)?[:\s\|]*([가-힣]{2,5})",
        ]:
            m = re.search(pat, text)
            if m and _ok(m.group(1)):
                return m.group(1).strip()
        return None

    @staticmethod
    def _extract_lessee(text: str) -> str | None:
        _NOT = {"임대인","임차인","성명","법인명"}
        for pat in [
            r"임\s*차\s*인\s*성\s*명[:\s]+([가-힣]{2,5})\s*(?:\(서명|\(서 명|$)",
            r"임\s*차\s*인\s*성\s*명[:\s]+([가-힣]{2,5})",
            r"임\s*차\s*인[:\s]*([가-힣]{2,5})\b",
        ]:
            m = re.search(pat, text)
            if m and m.group(1).strip() not in _NOT:
                return m.group(1).strip()
        return None

    @staticmethod
    def _extract_address(text: str) -> str | None:
        for pat in [
            r"주\s*택\s*소\s*재\s*지\s+([^\n]{10,80})",
            r"소\s*재\s*지[:\s]+([^\n]{5,80})",
            r"임\s*대\s*목\s*적\s*물[:\s]+([^\n]{5,80})",
        ]:
            m = re.search(pat, text)
            if m:
                addr = m.group(1).strip()
                if any(k in addr for k in ["시","구","동","로","길"]):
                    return addr
        return None

    @staticmethod
    def _extract_deposit(text: str) -> int | None:
        m = re.search(r"전체\s*(\d+)\s*억", text)
        if m:
            return int(m.group(1)) * 10000
        m = re.search(r"전체\s*([\d,]+)\s*억?\s*중", text)
        if m:
            a = int(m.group(1).replace(",", ""))
            return a * 10000 if a <= 100 else a // 10000
        cands = []
        for pat in [r"₩\s*([\d,]+)", r"#\s*([\d,]{9,15})", r"\(\s*[₩#]?\s*([\d,]{9,15})\s*\)"]:
            for fm in re.finditer(pat, text):
                try:
                    a = int(fm.group(1).replace(",", ""))
                    if 50_000_000 <= a <= 2_000_000_000:
                        cands.append(a // 10000)
                except ValueError:
                    pass
        if cands:
            return max(cands)
        for pat in [
            r"임\s*대\s*보\s*증\s*금[:\s]*금?\s*([\d,]+)\s*원",
            r"보\s*증\s*금[:\s]*금?\s*([\d,]+)\s*원",
            r"전세금[:\s]*금?\s*([\d,]+)\s*원",
        ]:
            m = re.search(pat, text)
            if m:
                try:
                    a = int(m.group(1).replace(",", ""))
                    return a // 10000 if a >= 10_000_000 else a
                except ValueError:
                    pass
        return None

    @staticmethod
    def _extract_monthly_rent(text: str) -> int | None:
        if re.search(r"영\s*원\s*정|₩\s*0\b|\(\s*₩0\s*\)", text):
            return 0
        for pat in [
            r"월\s*세[:\s]*금?\s*([\d,]+)\s*원",
            r"차\s*임[:\s]*금?\s*([\d,]+)\s*원",
            r"월\s*임\s*대\s*료[:\s]*금?\s*([\d,]+)\s*원",
        ]:
            m = re.search(pat, text)
            if m:
                try:
                    a = int(m.group(1).replace(",", ""))
                    return a // 10000 if a >= 1_000_000 else a
                except ValueError:
                    pass
        return 0

    @staticmethod
    def _extract_date(text: str, which: str) -> str | None:
        DP = r"(\d{4})[.\s년]\s*(\d{1,2})[.\s월]\s*(\d{1,2})"

        # 제외 날짜 수집
        excl: set[tuple] = set()
        for ep in [r"계\s*약\s*일[:\s]*" + DP, r"설\s*정\s*일\s*자[:\s]*" + DP]:
            for em in re.finditer(ep, text):
                excl.add((em.group(1), em.group(2).zfill(2), em.group(3).zfill(2)))
        for block_pat in [
            r"임\s*대\s*의\s*무\s*기\s*간([\s\S]{0,300}?)(?=\n\n|\Z)",
            r"입\s*주\s*(?:지\s*정\s*)?기\s*간?([\s\S]{0,150}?)(?=\n\n|\Z)",
        ]:
            bm = re.search(block_pat, text)
            if bm:
                for y, mo, d in re.findall(DP, bm.group(0)):
                    excl.add((y, mo.zfill(2), d.zfill(2)))

        # ① 임대차계약기간 키워드 + 범위 (최우선)
        pm = re.search(r"임\s*대\s*차\s*계\s*약\s*기\s*간([\s\S]{0,250})", text)
        if pm:
            rm = re.search(DP + r"\s*[.일]?\s*[~\-~]\s*" + DP, pm.group(0))
            if rm:
                if which == "start":
                    return f"{rm.group(1)}-{rm.group(2).zfill(2)}-{rm.group(3).zfill(2)}"
                return f"{rm.group(4)}-{rm.group(5).zfill(2)}-{rm.group(6).zfill(2)}"
            found = [(y, mo, d) for y, mo, d in re.findall(DP, pm.group(0))
                     if (y, mo.zfill(2), d.zfill(2)) not in excl]
            if found:
                y, mo, d = found[0] if which == "start" else found[-1]
                return f"{y}-{mo.zfill(2)}-{d.zfill(2)}"

        # ② pdfplumber 표 행
        tp = re.search(r"(?:계약기간|임대기간|임대차기간)[^\|]*\|\s*" + DP, text)
        if tp and which == "start":
            return f"{tp.group(1)}-{tp.group(2).zfill(2)}-{tp.group(3).zfill(2)}"

        # ③ fallback: 2020년 이후 날짜 중 excl 제외 후 첫/마지막
        dates = [(y, mo, d) for y, mo, d in re.findall(DP, text)
                 if (y, mo.zfill(2), d.zfill(2)) not in excl and int(y) >= 2020]
        if dates:
            y, mo, d = dates[0] if which == "start" else dates[-1]
            return f"{y}-{mo.zfill(2)}-{d.zfill(2)}"
        return None

    @staticmethod
    def _extract_contract_end(text: str, contract_start: str | None) -> str | None:
        DP = r"(\d{4})[.\s년]\s*(\d{1,2})[.\s월]\s*(\d{1,2})"
        excl: set[tuple] = set()
        for bp in [
            r"입\s*주\s*(?:지\s*정\s*)?기\s*간?([\s\S]{0,150}?)(?=\n\n|\Z)",
            r"임\s*대\s*의\s*무\s*기\s*간([\s\S]{0,300}?)(?=\n\n|\Z)",
        ]:
            bm = re.search(bp, text)
            if bm:
                for y, mo, d in re.findall(DP, bm.group(0)):
                    excl.add((y, mo.zfill(2), d.zfill(2)))

        pm = re.search(r"임\s*대\s*차\s*계\s*약\s*기\s*간([\s\S]{0,250})", text)
        if pm:
            rm = re.search(DP + r"\s*[.일]?\s*[~\-~]\s*" + DP, pm.group(0))
            if rm:
                return f"{rm.group(4)}-{rm.group(5).zfill(2)}-{rm.group(6).zfill(2)}"

        from datetime import date as _dt
        cands = []
        for y, mo, d in re.findall(DP, text):
            if (y, mo.zfill(2), d.zfill(2)) in excl:
                continue
            try:
                dt = _dt(int(y), int(mo), int(d))
                if contract_start:
                    sd = _dt.fromisoformat(contract_start)
                    if (dt - sd).days >= 365:
                        cands.append((dt, f"{y}-{mo.zfill(2)}-{d.zfill(2)}"))
                elif int(y) >= 2020:
                    cands.append((dt, f"{y}-{mo.zfill(2)}-{d.zfill(2)}"))
            except (ValueError, TypeError):
                pass
        return sorted(cands)[0][1] if cands else None

    @staticmethod
    def _find_special_terms_text(text: str) -> str | None:
        """다양한 계약서 구조에서 실제 특약 조항 텍스트를 찾아 반환합니다."""

        # ⓪ docx에서 직접 추출된 [특약사항] 태그 (가장 정확)
        if "\n\n[특약사항]\n" in text:
            _, st_raw = text.split("\n\n[특약사항]\n", 1)
            return st_raw.strip()

        # ① 민간임대 표준계약서: 임차인 서명란 → 번호 조항 → 개인정보 동의
        sig_m = re.search(
            r"임차인\s+성명\s*:.{0,60}\n(1\.\s+[\s\S]+?)(?=\n본인의\s*개인정보|\n\n본인|\Z)",
            text)
        if sig_m:
            block = sig_m.group(1).strip()
            if len(block) > 30 and not block.startswith("개인정보"):
                return block

        # ② 일반/표준 계약서: [특약사항] 괄호 뒤 내용
        bracket_m = re.search(
            r"[【\[《<＜]\s*특\s*약\s*사\s*항\s*[】\]》>＞]([\s\S]{10,2500})(?=\Z)",
            text, re.IGNORECASE)
        if bracket_m:
            block = bracket_m.group(1).strip()
            pi_idx = re.search(r"\d+\.\s*개인정보", block)
            if pi_idx:
                after_pi = block[pi_idx.start():]
                numbered = re.search(
                    r"(?:^|\n)(1\.\s+[\s\S]+?)(?=\n본인의\s*개인정보|\Z)",
                    after_pi)
                if numbered and len(numbered.group(1).strip()) > 30:
                    return numbered.group(1).strip()
            elif len(block) > 30:
                return block

        # ③ 키워드 fallback
        kw_m = re.search(
            r"특\s*약\s*사\s*항\s*[:\s]*([\s\S]{10,1000}?)(?=\n\n|\Z|개인정보)",
            text, re.IGNORECASE)
        if kw_m:
            block = kw_m.group(1).strip()
            if block and "\u300b" not in block[:3]:
                return block
        return None

        return ContractParser._find_special_terms_text(text)

    @classmethod
    def extract_risk_keywords(cls, text: str) -> list[str]:
        return list(dict.fromkeys(kw for kw in RISK_KEYWORDS if kw in text))

    @classmethod
    def extract_summary_keywords(cls, info: ContractInfo) -> list[str]:
        kws = []
        if info.deposit_amount:
            kws.append(f"보증금 {info.deposit_amount}만원")
        if info.address:
            kws.extend(re.findall(r"[가-힣]{2,4}[동구]", info.address))
        if info.special_terms:
            kws.extend(cls.extract_risk_keywords(info.special_terms))
        kws.extend(["전세계약", "위험 진단", "임차인 보호"])
        return kws
