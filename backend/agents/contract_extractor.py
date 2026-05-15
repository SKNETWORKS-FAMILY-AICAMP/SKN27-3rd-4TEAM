"""
contract_extractor — 임대차계약서에서 정보를 구조화 추출하는 에이전트

지원 형식: PDF, DOCX

추출 필드:
  - 주소, 전세금, 계약기간, 전용면적 (기본정보)
  - 특약 조항 목록 (선택)

판정 로직:
  - 월세 계약 감지 → "순수 전세만 분석 가능" 리턴 + 종료
  - 기본정보(주소, 전세금) 누락 → "분석불가" 리턴 + 종료
  - 특약은 있으면 추출, 없으면 빈 리스트
"""

import json
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from io import BytesIO
from pydantic import BaseModel, Field
from langchain_core.prompts import ChatPromptTemplate
from backend.config import get_llm


# ── 데이터 모델 ──────────────────────────────────────────

class ContractData(BaseModel):
    address: str | None = Field(None, description="매물 주소")
    deposit: int | None = Field(None, description="전세금 (만원)")
    contract_period: str | None = Field(None, description="계약기간 (예: 2025.03~2027.03)")
    area_m2: float | None = Field(None, description="전용면적 (㎡)")
    special_terms: list[str] = Field(default_factory=list, description="특약 조항 목록")
    has_monthly_rent: bool = Field(False, description="월세 포함 여부")


class ExtractionResult(BaseModel):
    success: bool
    data: ContractData | None = None
    message: str = ""


# ── 텍스트 추출 ──────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, str | None]:
    """파일에서 텍스트 추출. (텍스트, 에러메시지) 반환"""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        try:
            text = extract_text_from_docx(file_bytes)
        except Exception:
            return "", "DOC 파일은 DOCX로 변환 후 업로드해 주세요. (파일 → 다른 이름으로 저장 → .docx)"
    else:
        return "", f"지원하지 않는 파일 형식입니다: .{ext} (PDF 또는 DOCX만 가능)"

    if not text:
        return "", "파일에서 텍스트를 추출할 수 없습니다. 파일을 확인해 주세요."

    return text, None


# ── LLM 프롬프트 ────────────────────────────────────────

EXTRACTION_PROMPT = ChatPromptTemplate.from_messages([
    ("system", """당신은 한국 임대차계약서 분석 전문가입니다.
주어진 계약서 텍스트에서 아래 정보를 정확히 추출하세요.

반드시 아래 JSON 형식으로만 응답하세요. 다른 텍스트는 포함하지 마세요.

{{
  "address": "매물 주소 (없으면 null)",
  "deposit": 전세금 만원 단위 숫자 (없으면 null),
  "contract_period": "계약기간 (없으면 null)",
  "area_m2": 전용면적 숫자 (없으면 null),
  "special_terms": ["특약1", "특약2"],
  "has_monthly_rent": 월세가 있으면 true, 순수전세면 false
}}

규칙:
- 전세금은 반드시 만원 단위 정수로 변환 (예: 2억5천만원 → 25000)
- 특약이 없으면 빈 배열 []
- 월세(월임대료) 금액이 0보다 크면 has_monthly_rent를 true로
- 월임대료가 0원이거나 "영원"이면 순수 전세이므로 has_monthly_rent를 false로
- 표준 계약서 양식에 "월임대료"라는 용어가 나와도, 실제 금액이 0원이면 false
- 확인할 수 없는 정보는 null로"""),
    ("human", "다음은 임대차계약서 텍스트입니다:\n\n{contract_text}")
])


# ── 핵심 로직 ────────────────────────────────────────────

def parse_llm_response(response_text: str) -> ContractData:
    text = response_text.strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
        text = text.rsplit("```", 1)[0]
    parsed = json.loads(text)
    return ContractData(**parsed)


def extract_contract(file_bytes: bytes, filename: str) -> ExtractionResult:
    """계약서 파일(PDF/DOCX)에서 정보 추출 + 검증"""

    # 1. 파일 → 텍스트
    text, error = extract_text(file_bytes, filename)
    if error:
        return ExtractionResult(success=False, message=error)

    # 2. LLM으로 구조화 추출
    llm = get_llm(temperature=0.0)
    chain = EXTRACTION_PROMPT | llm
    response = chain.invoke({"contract_text": text[:12000]})

    try:
        data = parse_llm_response(response.content)
    except (json.JSONDecodeError, Exception):
        return ExtractionResult(
            success=False,
            message="계약서 분석 중 오류가 발생했습니다. 다시 시도해 주세요."
        )

    # 3. 월세 계약 체크
    if data.has_monthly_rent:
        return ExtractionResult(
            success=False,
            message="순수 전세 계약만 분석 가능합니다. 월세가 포함된 계약서는 지원하지 않습니다."
        )

    # 4. 기본정보 검증
    if not data.address or not data.deposit:
        return ExtractionResult(
            success=False,
            message="분석불가: 계약서에서 주소 또는 전세금 정보를 확인할 수 없습니다. 계약서를 다시 확인해 주세요."
        )

    # 5. 성공
    return ExtractionResult(
        success=True,
        data=data,
        message=f"계약서 분석 완료: {data.address} / 전세금 {data.deposit:,}만원"
        + (f" / 특약 {len(data.special_terms)}건" if data.special_terms else "")
    )


# ── 직접 입력 (파일 없이) ────────────────────────────────

def create_manual_input(address: str, deposit: int,
                        area_m2: float | None = None,
                        contract_period: str | None = None) -> ExtractionResult:
    """사용자가 파일 없이 직접 입력한 경우"""
    if not address or not deposit:
        return ExtractionResult(
            success=False,
            message="분석불가: 주소와 전세금은 필수 입력입니다."
        )

    data = ContractData(
        address=address,
        deposit=deposit,
        area_m2=area_m2,
        contract_period=contract_period,
        special_terms=[],
        has_monthly_rent=False,
    )

    return ExtractionResult(
        success=True,
        data=data,
        message=f"입력 완료: {address} / 전세금 {deposit:,}만원"
    )
