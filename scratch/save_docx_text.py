import docx
import os

def save_docx_text(file_name):
    file_path = os.path.join(r"c:\dev\project\SKN27-3rd-4TEAM\docs", file_name)
    output_path = os.path.join(r"c:\dev\project\SKN27-3rd-4TEAM\scratch", f"{file_name}.txt")
    
    doc = docx.Document(file_path)
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            text.append(" | ".join(c.text.strip() for c in row.cells))
            
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print(f"Saved {file_name} text to {output_path}")

if __name__ == "__main__":
    save_docx_text("가상계약서.docx")
    save_docx_text("평창동2.docx")
