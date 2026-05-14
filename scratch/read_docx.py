import docx
import os

def print_docx_content(file_path):
    print(f"\n--- Content of {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    for table in doc.tables:
        for row in table.rows:
            print(" | ".join(c.text.strip() for c in row.cells))

if __name__ == "__main__":
    print_docx_content(r"c:\dev\project\SKN27-3rd-4TEAM\docs\가상계약서.docx")
    print_docx_content(r"c:\dev\project\SKN27-3rd-4TEAM\docs\평창동2.docx")
